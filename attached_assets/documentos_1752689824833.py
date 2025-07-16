import os
import re
from docx import Document
from datetime import datetime

def extrair_slug_funcionalidade(texto):
    texto = texto.lower()
    palavras_chave = re.findall(
        r"(tela|cadastro|consulta|estoque|venda|cancelamento|painel|dashboard|relatório|validar|token|produto|pedido|cliente|usuário|reenvio|sms|cnpj)",
        texto
    )
    if palavras_chave:
        slug = "_".join(dict.fromkeys(palavras_chave))
    else:
        slug = re.sub(r"[^\w\s-]", "", texto)
        slug = re.sub(r"\s+", "_", slug)
        slug = slug[:60]
    return slug or "demanda_funcional"

def gerar_nome_arquivo(nome_base, prefixo):
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return f"{prefixo} - {nome_base}_{timestamp}.docx"

def salvar_prd(conteudo, demanda_funcional):
    try:
        nome_base = extrair_slug_funcionalidade(demanda_funcional)
        pasta = os.path.join("documentos")
        os.makedirs(pasta, exist_ok=True)
        caminho = os.path.join(pasta, gerar_nome_arquivo(nome_base, "PRD"))

        doc = Document()
        doc.add_heading("Product Requirements Document", 0)
        for linha in conteudo.strip().split("\n"):
            doc.add_paragraph(linha.strip())
        doc.save(caminho)

        print(f"✅ PRD salvo em: {caminho}")
    except Exception as e:
        print(f"❌ Erro ao salvar PRD: {e}")

def salvar_tasks(conteudo, demanda_funcional):
    try:
        nome_base = extrair_slug_funcionalidade(demanda_funcional)
        pasta = os.path.join("documentos")
        os.makedirs(pasta, exist_ok=True)
        caminho = os.path.join(pasta, gerar_nome_arquivo(nome_base, "Tasks"))

        doc = Document()
        doc.add_heading("Tasks (User Stories)", 0)
        for linha in conteudo.strip().split("\n"):
            doc.add_paragraph(linha.strip())
        doc.save(caminho)

        print(f"✅ Tasks salvas em: {caminho}")
    except Exception as e:
        print(f"❌ Erro ao salvar tasks: {e}")

